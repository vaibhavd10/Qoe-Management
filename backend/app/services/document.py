import os
import uuid
from typing import Optional, List, Dict, Any
from sqlalchemy.orm import Session
from datetime import datetime
from fastapi import HTTPException, status, UploadFile
import pandas as pd
import magic
from pathlib import Path

from app.models.document import Document, DocumentStatus, DocumentType
from app.models.project import Project
from app.schemas.document import DocumentCreate, DocumentUpdate
from app.core.config import settings
from app.tasks.document_processing import process_document_task


class DocumentService:
    """Service for handling document operations"""
    
    def __init__(self, db: Session):
        self.db = db
    
    def get_by_id(self, document_id: int) -> Optional[Document]:
        """Get document by ID"""
        return self.db.query(Document).filter(Document.id == document_id).first()
    
    def get_all(self, skip: int = 0, limit: int = 100, project_id: Optional[int] = None) -> List[Document]:
        """Get all documents with pagination and optional project filtering"""
        query = self.db.query(Document)
        
        if project_id:
            query = query.filter(Document.project_id == project_id)
        
        return query.offset(skip).limit(limit).all()
    
    def get_count(self, project_id: Optional[int] = None) -> int:
        """Get total document count"""
        query = self.db.query(Document)
        
        if project_id:
            query = query.filter(Document.project_id == project_id)
        
        return query.count()
    
    def upload_document(self, file: UploadFile, project_id: int, user_id: int) -> Document:
        """Upload and process a document"""
        # Validate project exists
        project = self.db.query(Project).filter(Project.id == project_id).first()
        if not project:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Project not found"
            )
        
        # Validate file size
        if file.size > settings.MAX_FILE_SIZE:
            raise HTTPException(
                status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
                detail="File too large"
            )
        
        # Validate file extension
        file_ext = Path(file.filename).suffix.lower()
        if file_ext not in settings.ALLOWED_EXTENSIONS:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"File type not allowed. Allowed types: {settings.ALLOWED_EXTENSIONS}"
            )
        
        # Generate unique filename
        unique_filename = f"{uuid.uuid4()}{file_ext}"
        file_path = os.path.join(settings.UPLOAD_DIR, unique_filename)
        
        # Save file
        try:
            with open(file_path, "wb") as buffer:
                content = file.file.read()
                buffer.write(content)
        except Exception as e:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Error saving file: {str(e)}"
            )
        
        # Get MIME type
        mime_type = magic.from_file(file_path, mime=True)
        
        # Create document record
        document = Document(
            filename=unique_filename,
            original_filename=file.filename,
            file_path=file_path,
            file_size=file.size,
            mime_type=mime_type,
            project_id=project_id,
            uploaded_by=user_id,
            status=DocumentStatus.PENDING
        )
        
        self.db.add(document)
        self.db.commit()
        self.db.refresh(document)
        
        # Queue processing task
        process_document_task.delay(document.id)
        
        return document
    
    def update(self, document_id: int, document_data: DocumentUpdate) -> Document:
        """Update document"""
        document = self.get_by_id(document_id)
        if not document:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Document not found"
            )
        
        # Update fields
        update_data = document_data.dict(exclude_unset=True)
        for field, value in update_data.items():
            if hasattr(document, field):
                setattr(document, field, value)
        
        document.updated_at = datetime.utcnow()
        
        self.db.commit()
        self.db.refresh(document)
        
        return document
    
    def delete(self, document_id: int) -> bool:
        """Delete document"""
        document = self.get_by_id(document_id)
        if not document:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Document not found"
            )
        
        # Delete file from filesystem
        try:
            if os.path.exists(document.file_path):
                os.remove(document.file_path)
        except Exception as e:
            # Log error but don't fail the deletion
            pass
        
        self.db.delete(document)
        self.db.commit()
        
        return True
    
    def get_by_project(self, project_id: int) -> List[Document]:
        """Get documents by project"""
        return self.db.query(Document).filter(Document.project_id == project_id).all()
    
    def get_by_status(self, status: DocumentStatus, project_id: Optional[int] = None) -> List[Document]:
        """Get documents by status"""
        query = self.db.query(Document).filter(Document.status == status)
        
        if project_id:
            query = query.filter(Document.project_id == project_id)
        
        return query.all()
    
    def get_by_type(self, document_type: DocumentType, project_id: Optional[int] = None) -> List[Document]:
        """Get documents by type"""
        query = self.db.query(Document).filter(Document.document_type == document_type)
        
        if project_id:
            query = query.filter(Document.project_id == project_id)
        
        return query.all()
    
    def classify_document(self, document_id: int) -> DocumentType:
        """Classify document type based on content"""
        document = self.get_by_id(document_id)
        if not document:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Document not found"
            )
        
        # Simple classification based on filename patterns
        filename_lower = document.original_filename.lower()
        
        classification_keywords = {
            DocumentType.GENERAL_LEDGER: ["gl", "general ledger", "ledger"],
            DocumentType.PROFIT_LOSS: ["pl", "p&l", "profit loss", "income statement"],
            DocumentType.BALANCE_SHEET: ["bs", "balance sheet", "statement of position"],
            DocumentType.TRIAL_BALANCE: ["tb", "trial balance", "trial"],
            DocumentType.PAYROLL: ["payroll", "payroll register", "salary"],
            DocumentType.CASH_FLOW: ["cash flow", "cf", "statement of cash flows"],
        }
        
        for doc_type, keywords in classification_keywords.items():
            if any(keyword in filename_lower for keyword in keywords):
                # Update document classification
                document.document_type = doc_type
                document.classification_confidence = 80
                self.db.commit()
                return doc_type
        
        # Default classification
        document.document_type = DocumentType.OTHER
        document.classification_confidence = 50
        self.db.commit()
        
        return DocumentType.OTHER
    
    def extract_data(self, document_id: int) -> Dict[str, Any]:
        """Extract data from document"""
        document = self.get_by_id(document_id)
        if not document:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Document not found"
            )
        
        try:
            if document.is_excel:
                return self._extract_excel_data(document)
            elif document.is_csv:
                return self._extract_csv_data(document)
            elif document.is_pdf:
                return self._extract_pdf_data(document)
            elif document.is_word:
                return self._extract_word_data(document)
            else:
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail="Unsupported file type for data extraction"
                )
        except Exception as e:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Error extracting data: {str(e)}"
            )
    
    def _extract_excel_data(self, document: Document) -> Dict[str, Any]:
        """Extract data from Excel file"""
        try:
            df = pd.read_excel(document.file_path, sheet_name=None)
            
            # If multiple sheets, process each
            if isinstance(df, dict):
                data = {}
                for sheet_name, sheet_df in df.items():
                    data[sheet_name] = {
                        "headers": sheet_df.columns.tolist(),
                        "data": sheet_df.to_dict('records')[:100],  # Limit to first 100 rows
                        "row_count": len(sheet_df),
                        "column_count": len(sheet_df.columns)
                    }
                
                # Update document metrics
                document.row_count = sum(sheet['row_count'] for sheet in data.values())
                document.column_count = max(sheet['column_count'] for sheet in data.values())
            else:
                data = {
                    "headers": df.columns.tolist(),
                    "data": df.to_dict('records')[:100],
                    "row_count": len(df),
                    "column_count": len(df.columns)
                }
                
                document.row_count = len(df)
                document.column_count = len(df.columns)
            
            document.extracted_data = data
            self.db.commit()
            
            return data
        except Exception as e:
            raise Exception(f"Error processing Excel file: {str(e)}")
    
    def _extract_csv_data(self, document: Document) -> Dict[str, Any]:
        """Extract data from CSV file"""
        try:
            df = pd.read_csv(document.file_path)
            
            data = {
                "headers": df.columns.tolist(),
                "data": df.to_dict('records')[:100],
                "row_count": len(df),
                "column_count": len(df.columns)
            }
            
            document.row_count = len(df)
            document.column_count = len(df.columns)
            document.extracted_data = data
            self.db.commit()
            
            return data
        except Exception as e:
            raise Exception(f"Error processing CSV file: {str(e)}")
    
    def _extract_pdf_data(self, document: Document) -> Dict[str, Any]:
        """Extract data from PDF file"""
        try:
            import pdfplumber
            
            text_content = []
            tables = []
            
            with pdfplumber.open(document.file_path) as pdf:
                for page in pdf.pages:
                    # Extract text
                    text = page.extract_text()
                    if text:
                        text_content.append(text)
                    
                    # Extract tables
                    page_tables = page.extract_tables()
                    if page_tables:
                        tables.extend(page_tables)
            
            data = {
                "text_content": text_content,
                "tables": tables,
                "page_count": len(text_content),
                "table_count": len(tables)
            }
            
            document.extracted_data = data
            self.db.commit()
            
            return data
        except Exception as e:
            raise Exception(f"Error processing PDF file: {str(e)}")
    
    def _extract_word_data(self, document: Document) -> Dict[str, Any]:
        """Extract data from Word document"""
        try:
            from docx import Document as DocxDocument
            
            doc = DocxDocument(document.file_path)
            
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text)
                    table_data.append(row_data)
                tables.append(table_data)
            
            data = {
                "paragraphs": paragraphs,
                "tables": tables,
                "paragraph_count": len(paragraphs),
                "table_count": len(tables)
            }
            
            document.extracted_data = data
            self.db.commit()
            
            return data
        except Exception as e:
            raise Exception(f"Error processing Word document: {str(e)}")
    
    def get_processing_status(self, document_id: int) -> Dict[str, Any]:
        """Get document processing status"""
        document = self.get_by_id(document_id)
        if not document:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Document not found"
            )
        
        progress = 0
        if document.status == DocumentStatus.PENDING:
            progress = 0
        elif document.status == DocumentStatus.PROCESSING:
            progress = 50
        elif document.status == DocumentStatus.COMPLETED:
            progress = 100
        elif document.status == DocumentStatus.ERROR:
            progress = 0
        
        return {
            "id": document.id,
            "filename": document.original_filename,
            "status": document.status,
            "progress": progress,
            "error_message": document.error_message,
            "processing_time": document.processing_time
        }